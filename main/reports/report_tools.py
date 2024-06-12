import json
import hashlib
import PyPDF2
from docx import Document
from openpyxl.drawing.image import Image as XlsxImage
from openpyxl.worksheet.page import PageMargins
from PIL import Image as PilImage
import os
import subprocess
import logging
import re
from data.models import ReportAuthenticity

# Creation of logger
logger = logging.getLogger(__name__)


def replace_variables_in_docx(docx_file_path, variables_dict):
    temporal_docx_file_path = 'media/file_templates/temporal.docx'
    doc = Document(docx_file_path)

    # Reemplazar variables en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for match in re.findall(r'\{\{.*?}}', run.text):
                key = match[2:-2]
                if key in variables_dict:
                    run.text = run.text.replace(match, variables_dict[key])

    # Reemplazar variables en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for match in re.findall(r'\{\{.*?}}', run.text):
                            key = match[2:-2]
                            if key in variables_dict:
                                run.text = run.text.replace(match, variables_dict[key])

    doc.save(temporal_docx_file_path)
    return temporal_docx_file_path


def keep_first_page(pdf_path):
    # Open the PDF file
    reader = PyPDF2.PdfReader(pdf_path)

    # Create a new PDF writer
    writer = PyPDF2.PdfWriter()

    # Add the first page to the writer
    writer.add_page(reader.pages[0])

    # Write the output to a new file
    with open(pdf_path, 'wb') as output_file:
        writer.write(output_file)


def cm_to_pixels(cm: float) -> int:
    """
    Convert centimeters to pixels.

    Args:
        cm: The number of centimeters to convert.

    Returns:
        The number of pixels equivalent to the provided centimeters.
    """
    dpi = 96
    return int(dpi * cm / 2.54)


def scale_image_from_height(image_path: str, desired_height_cm: float):
    # Load the image with PIL to get its size
    with PilImage.open(image_path) as pil_img:
        original_width, original_height = pil_img.size

    # Convert the height from cm to pixels and set the height of the image
    dpi = 96
    desired_height_px = cm_to_pixels(desired_height_cm)

    # Calculate the new width to maintain aspect ratio
    scale_factor = desired_height_px / original_height
    desired_width_px = int(original_width * scale_factor)

    return desired_width_px, desired_height_px


def add_image_to_worksheet(image_path: str, cell: str, activesheet, width: int, height: int):
    # Load and add the image with openpyxl
    img = XlsxImage(image_path)
    img.width = width
    img.height = height

    # Add image to the specified cell
    activesheet.add_image(img, cell)


def convert_to_pdf(temporary_file_path: str, pdf_path: str):
    convertion_command = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        os.path.dirname(pdf_path),
        temporary_file_path
    ]
    process = subprocess.Popen(convertion_command)
    process.wait()  # Wait for the process to finish before returning
    return pdf_path if process.returncode == 0 else None  # Check return code for success


def merge_pdf_files(existing_pdf_path: str, added_pdf_path: str, output_pdf_path: str):
    merger = PyPDF2.PdfMerger()
    merger.append(PyPDF2.PdfReader(open(existing_pdf_path, 'rb')))
    merger.append(PyPDF2.PdfReader(open(added_pdf_path, 'rb')))
    merger.write(open(output_pdf_path, 'wb'))
    return output_pdf_path


def xlsx_sheet_presets(sheet):
    sheet.page_setup.paperSize = sheet.PAPERSIZE_LETTER
    sheet.page_margins = PageMargins(
        top=0.2,
        left=0.25,
        right=0.01,
        header=0.1,
        footer=0.1,
        bottom=0.2
    )


def generate_and_save_authenticity_chain(data, report_name):
    # Create a copy of the data dictionary and remove 'cadena_autenticidad'
    data_without_authenticity = data.copy()
    data_without_authenticity.pop('cadena_autenticidad', None)

    # Convert the modified data dictionary to a JSON string
    content = json.dumps(data_without_authenticity)

    # Generate a checksum of the content
    content_checksum = hashlib.sha256(content.encode()).hexdigest()

    # Check if a report with the same name and content checksum already exists
    existing_report = ReportAuthenticity.objects.filter(report_name=report_name, content=content_checksum).first()

    if existing_report:
        # If the report exists and the content hasn't changed, use the existing authenticity chain
        authenticity_chain = existing_report.authenticity_chain
    else:
        # If the report doesn't exist or the content has changed, generate a new authenticity chain
        authenticity_chain = hashlib.sha256(content.encode()).hexdigest()[:50]

        # Save the new authenticity chain and content checksum to the database
        ReportAuthenticity.objects.update_or_create(
            report_name=report_name,
            defaults={
                'authenticity_chain': authenticity_chain,
                'content': content_checksum
            },
        )
    return authenticity_chain

